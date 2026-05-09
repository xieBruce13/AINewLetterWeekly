"""
export_product_docx.py
======================
Generates docs/ZenoNews_Product_Brief.docx from the full PRD content
matching docs/Newsletter机制说明.md (8 chapters).

Usage:
    python tools/export_product_docx.py
"""

from pathlib import Path
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement


# ── helpers ──────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_table(doc, headers: list[str], rows: list[list[str]], col_widths: list[float] | None = None):
    """Add a styled table with header row."""
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"

    # Header row
    hdr = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr.cells[i]
        set_cell_bg(cell, "2D2D2D")
        run = cell.paragraphs[0].add_run(h)
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(9)

    # Data rows
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "F9F9F9" if ri % 2 == 0 else "FFFFFF"
        for ci, cell_text in enumerate(row):
            cell = tr.cells[ci]
            set_cell_bg(cell, bg)
            p = cell.paragraphs[0]
            run = p.add_run(str(cell_text))
            run.font.size = Pt(9)
            if ci == 0:
                run.bold = True

    # Optional column widths
    if col_widths:
        for ci, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[ci].width = Inches(w)

    doc.add_paragraph()  # spacing after table
    return table


def h1(doc, text: str):
    p = doc.add_heading(text, level=1)
    p.runs[0].font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)


def h2(doc, text: str):
    p = doc.add_heading(text, level=2)
    p.runs[0].font.color.rgb = RGBColor(0xE8, 0x5D, 0x3A)


def h3(doc, text: str):
    doc.add_heading(text, level=3)


def h4(doc, text: str):
    doc.add_heading(text, level=4)


def body(doc, text: str):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(10) if p.runs else None


def bullet(doc, text: str):
    doc.add_paragraph(text, style="List Bullet")


def note(doc, text: str):
    p = doc.add_paragraph(text)
    p.paragraph_format.left_indent = Inches(0.3)
    for run in p.runs:
        run.font.italic = True
        run.font.size = Pt(9)
        run.font.color.rgb = RGBColor(0x55, 0x55, 0x55)


def divider(doc):
    doc.add_paragraph("─" * 60)


# ── main ─────────────────────────────────────────────────────────────────────

def build():
    doc = Document()

    # ── Page margins ──
    for section in doc.sections:
        section.top_margin = Inches(1)
        section.bottom_margin = Inches(1)
        section.left_margin = Inches(1.2)
        section.right_margin = Inches(1.2)

    # ── Cover ──
    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title_p.add_run("ZenoNews — AI 行业周报产品说明（PRD）")
    run.bold = True
    run.font.size = Pt(22)
    run.font.color.rgb = RGBColor(0xE8, 0x5D, 0x3A)

    sub_p = doc.add_paragraph()
    sub_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    sub_run = sub_p.add_run(
        "编辑机制 · 网站个性化 · 信源范围 · OpenAI 成本测算"
    )
    sub_run.font.size = Pt(12)
    sub_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ══════════════════════════════════════════════════════════════
    # 1 文档控制
    # ══════════════════════════════════════════════════════════════
    h1(doc, "1  文档控制")
    add_table(
        doc,
        ["项", "内容"],
        [
            ["文档目的", "对齐编辑流水线、网站个性化、信源范围与模型/成本假设。"],
            ["适用范围", "仓库内 skill/ SOP、.claude/agents/、tools/ 抓取管线、web/ 阅读站。"],
            [
                "权威与冲突处理",
                "成稿结构与编辑规则以 skill/SKILL.md 为准；RSS 等可配置抓取清单以 tools/sources.yaml 为准。正文信源表与 YAML 同步维护。",
            ],
        ],
        [2.0, 5.0],
    )

    # ══════════════════════════════════════════════════════════════
    # 2 编辑与发布机制
    # ══════════════════════════════════════════════════════════════
    h1(doc, "2  编辑与发布机制")

    h2(doc, "2.1  周报与编辑部（Zeno）")
    body(doc, "默认范围（可被 Step 0 覆盖）")
    add_table(
        doc,
        ["参数", "默认值"],
        [
            ["时间窗", "最近 7 天（截止「当前运行时刻」）"],
            ["模块", "models + products"],
            ["语言", "与用户请求一致；中文请求默认中文成稿"],
            ["主条目数量", "每模块约 top 3（密集周可放宽至 5，见 SKILL）"],
        ],
        [2.5, 4.5],
    )

    body(doc, "成稿结构（固定顺序）")
    body(
        doc,
        "本周结论 → 模型模块 → 产品模块 → 本期初创聚焦 → 简讯（两列表格）→ 编辑部判断 → 参考来源。篇幅目标 6–8 页 PDF。",
    )
    doc.add_paragraph()

    body(doc, "取舍规则（摘录）")
    bullet(doc, "模型主条：评分阈值与维度见 skill/rubric.json；主文门槛含「可核对的一手/技术源」。")
    bullet(doc, "产品主条：需用户可感知发布 + 多源印证（见 SKILL Gate）。")
    bullet(doc, "结构性：公司多样性上限；每期至少 1 条非八大厂初创/独立产品（定义见 SKILL）。")
    doc.add_paragraph()

    h2(doc, "2.2  网站个性化")
    add_table(
        doc,
        ["步骤", "行为", "数据来源"],
        [
            ["画像", "角色、公司、当前项目、focus_topics[]、dislikes[]", "user_profiles"],
            ["候选池", "标签重叠、向量相似、跨期刊、LOOKBACK_ISSUES = 4", "news_items、embedding"],
            [
                "精排文案",
                "为 Top N 生成 personalized blurb / reason；缺 OPENAI_API_KEY 时降级",
                "web/lib/personalization/rerank.ts",
            ],
            ["聊天", "流式回复 + 可选新闻检索工具", "web/app/api/chat/route.ts"],
            ["长期记忆", "从对话抽取事实写入记忆表", "web/lib/agent/memory.ts"],
        ],
        [1.5, 3.0, 2.5],
    )

    # ══════════════════════════════════════════════════════════════
    # 3 更新频率与时间范围
    # ══════════════════════════════════════════════════════════════
    h1(doc, "3  更新频率与时间范围")
    add_table(
        doc,
        ["项", "定义"],
        [
            ["周报素材窗", "默认 7 天；每期 Step 0 写入 run_header 可改。"],
            ["个性化回看", "约 4 期 × 每期刊内条目，用于跨期对比（非「只显示第 4 期」）。"],
            ["入库新鲜度", "由 P2 抓取频率与 sync_to_db 运行时刻决定；无内置 cron 强制日更。"],
            ["站外 digest", "配置为每周一次触发（与部署平台上 cron 一致）。"],
        ],
        [2.5, 4.5],
    )

    # ══════════════════════════════════════════════════════════════
    # 4 数据来源
    # ══════════════════════════════════════════════════════════════
    h1(doc, "4  数据来源")
    body(
        doc,
        "分四块：RSS/Atom（YAML）、Reddit、Hacker News、深度刊补充（Newsletter / YC / Product Hunt 等）；配图单独一小节。",
    )

    # 4.1 RSS/Atom
    h2(doc, "4.1  RSS/Atom（tools/sources.yaml）")
    body(
        doc,
        "图例 — tier：official 官方站点 / press 科技媒体；module：管线侧默认归类，实际仍以正文分类为准。",
    )

    h3(doc, "A1 官方 — 模型实验室与基础设施（20）")
    add_table(
        doc,
        ["名称", "Feed URL"],
        [
            ["OpenAI Blog", "openai.com/blog/rss/"],
            ["Anthropic News", "anthropic.com/news/rss.xml"],
            ["Google DeepMind", "deepmind.google/blog/feed/basic/"],
            ["Google AI Blog", "blog.google/technology/ai/rss/"],
            ["Meta AI Blog", "ai.meta.com/blog/rss/"],
            ["Mistral AI News", "mistral.ai/news/rss.xml"],
            ["HuggingFace Blog", "huggingface.co/blog/feed.xml"],
            ["Cohere Blog", "cohere.com/blog/rss"],
            ["xAI News", "x.ai/news/rss"],
            ["Stability AI Blog", "stability.ai/news/rss.xml"],
            ["Together AI Blog", "together.ai/blog/rss.xml"],
            ["Groq Blog", "groq.com/blog/rss.xml"],
            ["Cerebras Blog", "cerebras.net/blog/rss.xml"],
            ["Perplexity Blog", "blog.perplexity.ai/rss"],
            ["Scale AI Blog", "scale.com/blog/rss.xml"],
            ["AWS Machine Learning Blog", "aws.amazon.com/blogs/machine-learning/feed/"],
            ["Microsoft AI Blog", "blogs.microsoft.com/ai/feed/"],
            ["NVIDIA AI Blog", "blogs.nvidia.com/feed/"],
            ["Allen Institute AI", "allenai.org/blog/rss.xml"],
            ["Qwen Blog (Alibaba)", "qwenlm.github.io/feed.xml"],
        ],
        [2.5, 4.5],
    )

    h3(doc, "A2 官方 — 创作 / 图像 / 视频 / 设计（13）")
    add_table(
        doc,
        ["名称", "Feed URL"],
        [
            ["Luma AI Blog", "lumalabs.ai/blog/rss.xml"],
            ["Runway Blog", "runwayml.com/blog/rss.xml"],
            ["Runway Research", "research.runwayml.com/feed.xml"],
            ["Pika Blog", "pika.art/blog/rss.xml"],
            ["Midjourney Updates", "updates.midjourney.com/rss.xml"],
            ["Adobe Blog AI", "blog.adobe.com/en/topics/ai-ml/feed"],
            ["Figma Blog", "figma.com/blog/rss.xml"],
            ["Canva Newsroom", "canva.com/newsroom/rss.xml"],
            ["ElevenLabs Blog", "elevenlabs.io/blog/rss.xml"],
            ["Suno Blog", "suno.com/blog/rss.xml"],
            ["Udio Blog", "udio.com/blog/rss.xml"],
            ["Krea AI Blog", "krea.ai/blog/rss.xml"],
            ["Kling AI Blog", "klingai.com/blog/rss.xml"],
        ],
        [2.5, 4.5],
    )

    h3(doc, "A3 官方 — 编程 / Agent / 开发者工具（10）")
    add_table(
        doc,
        ["名称", "Feed URL"],
        [
            ["Cursor Changelog", "cursor.com/changelog/rss.xml"],
            ["GitHub Blog AI", "github.blog/ai-and-ml/feed/"],
            ["Vercel Blog", "vercel.com/blog/rss.xml"],
            ["LangChain Blog", "blog.langchain.dev/rss/"],
            ["LlamaIndex Blog", "llamaindex.ai/blog/rss.xml"],
            ["Replit Blog", "blog.replit.com/rss"],
            ["Codeium (Windsurf) Blog", "codeium.com/blog/rss.xml"],
            ["Linear Blog", "linear.app/blog/rss.xml"],
            ["Notion Blog", "notion.com/blog/rss.xml"],
            ["Weights & Biases Blog", "wandb.ai/fully-connected/rss.xml"],
        ],
        [2.5, 4.5],
    )

    h3(doc, "A4 科技媒体 — press（12）")
    add_table(
        doc,
        ["名称", "Feed URL"],
        [
            ["TechCrunch AI", "techcrunch.com/category/artificial-intelligence/feed/"],
            ["VentureBeat AI", "venturebeat.com/ai/feed/"],
            ["The Verge AI", "theverge.com/ai-artificial-intelligence/rss/index.xml"],
            ["Ars Technica Technology", "feeds.arstechnica.com/arstechnica/technology-lab"],
            ["Wired AI", "wired.com/feed/tag/ai/latest/rss"],
            ["MIT Technology Review", "technologyreview.com/feed/"],
            ["9to5Mac AI", "9to5mac.com/guides/artificial-intelligence/feed/"],
            ["The Information", "theinformation.com/feed"],
            ["CNBC Technology", "cnbc.com/id/19854910/device/rss/rss.html"],
            ["Fortune Technology", "fortune.com/section/technology/feed/"],
            ["ZDNet AI", "zdnet.com/topic/artificial-intelligence/rss.xml"],
            ["IEEE Spectrum AI", "spectrum.ieee.org/feeds/blog/artificial-intelligence.rss"],
        ],
        [2.5, 4.5],
    )
    note(doc, "RSS 小计（当前 active）— 官方 42 + 媒体 12 = 54 条 Feed。")

    # 4.2 Reddit
    h2(doc, "4.2  Reddit（reddit_sources）")
    add_table(
        doc,
        ["r/", "模块倾向", "备注"],
        [
            ["MachineLearning", "model", "论文与研究讨论"],
            ["LocalLLaMA", "model", "开源权重与本地推理"],
            ["singularity", "—", "综合"],
            ["artificial", "—", "综合"],
            ["ChatGPT", "product", ""],
            ["ClaudeAI", "product", ""],
            ["StableDiffusion", "product", "图像"],
            ["midjourney", "product", ""],
            ["aivideo", "product", ""],
            ["AIArt", "product", ""],
            ["ChatGPTPromptEngineering", "product", "工作流/提示"],
        ],
        [2.0, 2.0, 3.0],
    )
    note(doc, "当前关闭（active: false）— Bing、SoftwareEngineering：保留在 YAML 中但不抓取。")

    # 4.3 Hacker News
    h2(doc, "4.3  Hacker News（hacker_news）")
    add_table(
        doc,
        ["参数", "值"],
        [
            ["启用", "enabled: true"],
            ["扫描条数", "前 max_stories: 500 条故事"],
            ["最低分数", "min_score: 50"],
            [
                "匹配方式",
                "标题命中关键词列表（含 gpt、claude、llm、openai、anthropic、gemini、diffusion、cursor、copilot、agent 等，全文见 tools/sources.yaml）",
            ],
        ],
        [2.0, 5.0],
    )

    # 4.4 Newsletter / YC / Product Hunt
    h2(doc, "4.4  Newsletter、YC、Product Hunt 与 SOP 补充")
    body(
        doc,
        "自动化抓取以 4.1–4.3 为准；下表为 SKILL 与行业常用信息源，由编辑或多智能体步骤按需浏览；未接入 RSS 时需手工补。",
    )

    h3(doc, "C1 — Newsletter（英文行业简报）")
    add_table(
        doc,
        ["名称", "侧重", "常见入口"],
        [
            ["TLDR AI", "日更 headline，覆盖面广", "tldr.tech"],
            ["Import AI", "周更，偏研究与产业解读", "Substack「Import AI」"],
            ["The Rundown AI", "综合 AI / 商业", "therundown.ai"],
            ["Ben's Bites", "产品、创业、小工具", "bensbites.co"],
            ["Superhuman AI", "工具上手与速递", "读者订阅信/官网"],
            ["AI Breakfast", "每日简报", "独立 newsletter"],
            ["AI Valley Newsletter", "工具与创业信号", "newsletter 订阅"],
            ["Last Week in AI", "周更汇总类", "独立/Substack"],
            ["AlphaSignal 等", "论文/模型速递", "按需订阅"],
        ],
        [2.0, 2.5, 2.5],
    )
    note(
        doc,
        "SKILL 中模型 Tier3 固定点名：Import AI、TLDR AI、The Rundown AI；产品 Tier3 还含 Ben's Bites、Superhuman AI、AI Breakfast、AI Valley 等。",
    )

    h3(doc, "C2 — Y Combinator、Product Hunt、工具聚合")
    add_table(
        doc,
        ["类型", "用途", "入口/用法"],
        [
            ["Y Combinator", "筛 AI 标签公司、批次名单、官网/changelog", "ycombinator.com/companies（行业筛 AI）"],
            ["Product Hunt", "当日/当周 AI 类 Launch、评论与热度", "producthunt.com · Topics: Artificial Intelligence"],
            ["Futurepedia", "工具库、分类检索", "futurepedia.io"],
            ["There's An AI For That (TAAFT)", "工具目录与趋势", "theresanaiforthat.com"],
            ["Toolify", "AI 工具榜与分类", "toolify.ai"],
            ["Indie Hackers", "独立开发者产品叙事", "indiehackers.com"],
            ["Google Trends / TikTok / YouTube / X", "分发与话题热度（只作侧证）", "各平台官方站点"],
        ],
        [2.0, 2.5, 2.5],
    )

    h3(doc, "C3 — SKILL 其余补充")
    bullet(doc, "模型：一手 Release Notes；HF Trending、GitHub Trending、arXiv、benchmark 仓库；Arena/LMSYS、Artificial Analysis、LiveBench 做交叉验证。")
    bullet(doc, "产品：各产品官方 Changelog（ChatGPT、Claude、Gemini、Perplexity、xAI 等，URL 随产品迭代）。")
    bullet(doc, "社区：X 上官方号/创始人、Reddit 扩展子版、Towards AI 社区、Hacker News（与 4.3 自动化关键词流互补）。")
    bullet(doc, "初创硬规则：每期至少 1 条非八大厂产品 — 常从 Product Hunt 日榜 AI、YC 目录、Indie Hackers、Ben's Bites、TAAFT 中兑现。")
    doc.add_paragraph()

    # 4.5 配图
    h2(doc, "4.5  配图（Unsplash 与 Logo）")
    add_table(
        doc,
        ["类型", "来源", "说明"],
        [
            [
                "卡片 / 列表主视觉",
                "Unsplash",
                "tools/fetch_unsplash_covers.py 调 Search API；图片来自 Unsplash CDN。环境变量：UNSPLASH_ACCESS_KEY。",
            ],
            [
                "品牌辨识",
                "Logo",
                "公司/产品标识以官方 Press Kit、品牌页或通用 Logo 收录服务为准；用于角标、列表图标或与 Unsplash 配图并存。",
            ],
            [
                "备选（控费外）",
                "OpenAI 图像 API",
                "tools/generate_cover_images.py 默认 gpt-image-1，仅按需启用；与 Unsplash 二选一。",
            ],
        ],
        [1.8, 1.5, 3.7],
    )

    # ══════════════════════════════════════════════════════════════
    # 5 模型配置
    # ══════════════════════════════════════════════════════════════
    h1(doc, "5  模型配置（代码常量）")
    add_table(
        doc,
        ["能力", "模型（仓库当前常量）", "文件"],
        [
            ["Enrichment 主调用", "gpt-4.1", "tools/ai_filter.py"],
            ["Fallback 轻量筛选", "gpt-4.1-mini", "tools/ai_filter.py"],
            ["入库向量", "text-embedding-3-small", "tools/sync_to_db.py、web/lib/embeddings.ts"],
            ["聊天", "gpt-4o-mini", "web/app/api/chat/route.ts"],
            ["个性化精排文案", "gpt-4o-mini", "web/lib/personalization/rerank.ts"],
            ["记忆抽取", "gpt-4o-mini", "web/lib/agent/memory.ts"],
            ["可选生图封面", "gpt-image-1（CLI 默认）", "tools/generate_cover_images.py"],
        ],
        [2.0, 2.5, 2.5],
    )
    note(
        doc,
        "多智能体 SOP（.claude/agents/*）不绑定单一 API 名；若在 Claude API 上跑整条链，计价以 Anthropic 控制台为准。",
    )

    # ══════════════════════════════════════════════════════════════
    # 6 OpenAI 成本测算
    # ══════════════════════════════════════════════════════════════
    h1(doc, "6  OpenAI 成本测算（P2 入库 + 网站）")
    body(
        doc,
        "以下假设计费均经 OpenAI API，模型与仓库当前常量一致：gpt-4.1（入库扩写）、text-embedding-3-small（向量）、gpt-4o-mini（首页精排、聊天、记忆）。单价以 OpenAI Pricing 为准，数字用于量级估算，实施前请用控制台用量校准。",
    )
    add_table(
        doc,
        ["模型", "测算单价（USD / 1M tokens）"],
        [
            ["GPT-4.1", "输入 $2.00 · 输出 $8.00"],
            ["GPT-4o-mini", "输入 $0.15 · 输出 $0.60"],
            ["text-embedding-3-small", "$0.02（仅输入侧计费）"],
        ],
        [3.0, 4.0],
    )

    body(doc, "固定情景参数（单线维持一条「典型产品站」）")
    add_table(
        doc,
        ["参数", "取值"],
        [
            ["月刊期数", "4（每周 1 期入库）"],
            ["每期 ai_filter.py 扩写条数", "30；每批 15 条 → 2 次/期 gpt-4.1"],
            ["单次 enrich 调用", "输入约 14k（长 system + 15 条摘要）· 输出约 35k（15 条长 schema）"],
            ["Embedding", "每期 30 条 × ~900 tokens ≈ 27k/期"],
            ["每活跃用户每月", "打开首页 4 次（周更）· 聊天 10 轮 · 记忆抽取 3 次"],
        ],
        [2.5, 4.5],
    )

    h2(doc, "6.1  基线：维持成本（与用户量无关）")
    body(doc, "仅 P2 管道：入库扩写 + 向量，不计任何用户请求。")
    add_table(
        doc,
        ["项", "Token 量级/月", "公式（USD）", "结果"],
        [
            ["GPT-4.1 入库", "入 112k · 出 280k", "112×2/1000 + 280×8/1000", "≈ $2.46"],
            ["Embedding", "入 ≈108k", "108×0.02/1000", "≈ $0.00"],
        ],
        [1.8, 1.8, 2.5, 1.0],
    )
    note(doc, "维持基数合计 ≈ $2.5 / 月（仅 OpenAI 文本与向量；不含 Unsplash、gpt-image-1、Vercel/域名等）。")

    h2(doc, "6.2  边际：每增加 1 名活跃用户（每月）")
    add_table(
        doc,
        ["行为", "单用户 token/月", "费用（USD）"],
        [
            ["首页精排 ×4", "入 20k · 出 8k", "20×0.15/1000 + 8×0.6/1000 ≈ $0.0078"],
            ["聊天 ×10 轮", "入 20k · 出 5k", "20×0.15/1000 + 5×0.6/1000 ≈ $0.0060"],
            ["记忆 ×3", "入 4.5k · 出 1.2k", "≈ $0.0014"],
        ],
        [2.0, 2.5, 2.5],
    )
    note(
        doc,
        "边际合计 ≈ $0.015 / 用户 / 月。总公式：月 OpenAI 文本/向量 ≈ $2.5 + N × $0.015。\n例：N = 50 → $3.3 / 月。",
    )

    h2(doc, "6.3  参考：日更入库、更强模型、50 名活跃（月度）")
    body(
        doc,
        "结论（预算口径）：在「日更入库、约 50 名活跃用户、对话与记忆用更强模型」的前提下，只要对首页精排做周更或缓存，OpenAI 文本与向量可压在约 ¥200–400 / 月（约 $28–55 / 月，按 ¥≈USD×7.2 量级换算，以控制台为准）。",
    )

    add_table(
        doc,
        ["模型", "测算（USD / 1M tokens）"],
        [
            ["GPT-4.1", "入 $2.00 · 出 $8.00"],
            ["GPT-4o", "入 $2.50 · 出 $10.00"],
            ["GPT-4o-mini", "入 $0.15 · 出 $0.60"],
            ["text-embedding-3-small", "$0.02（仅输入）"],
        ],
        [3.0, 4.0],
    )

    h3(doc, "推荐方案（落入 ¥200–400）：日更入库 + 首页 mini 或周更 4o + 对话记忆 4o")
    add_table(
        doc,
        ["分项", "模型与频次", "金额（USD/月）"],
        [
            ["A. 入库", "gpt-4.1，日更", "≈ $18.48"],
            ["B. Embedding", "text-embedding-3-small", "≈ $0.02"],
            [
                "C. 首页精排",
                "gpt-4o-mini，50×30 次/月（日更展示可依赖缓存与 cheap rerank，精排不必日日走贵模型）",
                "≈ $2.93",
            ],
            ["D. 聊天", "gpt-4o", "≈ $5.63"],
            ["E. 记忆", "gpt-4o", "≈ $1.16"],
        ],
        [1.2, 4.0, 1.8],
    )
    note(
        doc,
        "小计 ≈ $28.2 / 月 → 约 ¥203（@7.2），落在区间偏下。\n"
        "若首页文案也用 gpt-4o 但改为每周约 1 次/人（其余用缓存），则小计 ≈ $31.8 → 约 ¥229；"
        "叠加偶发重跑等，上沿可接近 ¥350–400，仍明显低于「每人每日 gpt-4o 精排」。",
    )

    h3(doc, "对照：未优化（每人每天 gpt-4o 精排）")
    add_table(
        doc,
        ["分项", "计算依据", "金额（USD/月）"],
        [
            ["A–B", "同推荐方案", "≈ $18.50"],
            ["C. 首页", "gpt-4o，50×30×（5k 入 + 2k 出）", "≈ $48.75"],
            ["D–E", "gpt-4o 聊天 + 记忆", "≈ $6.79"],
        ],
        [1.2, 4.0, 1.8],
    )
    note(
        doc,
        "合计 ≈ $74 / 月（约 ¥530+）— 账单大头在首页精排，不建议作为默认产品形态。",
    )

    h2(doc, "6.4  配图与非文本 API")
    add_table(
        doc,
        ["项", "说明"],
        [
            ["Unsplash", "通常免费层；$0 计入上表。"],
            ["gpt-image-1", "批量 AI 封面时另计，见图像定价页。"],
        ],
        [2.0, 5.0],
    )

    # ══════════════════════════════════════════════════════════════
    # 7 风险与待确认
    # ══════════════════════════════════════════════════════════════
    h1(doc, "7  风险与待确认")
    add_table(
        doc,
        ["ID", "风险", "缓解"],
        [
            ["R1", "Feed 失效或 403", "sources.yaml 单源 active: false；SKILL 人工补源"],
            ["R2", "日更 × gpt-4o 使首页调用量主导账单", "缓存精排结果、降频、或高峰用 mini"],
            ["R3", "标价与汇率变动", "以 OpenAI 控制台 30 天用量为准"],
        ],
        [0.6, 3.0, 3.4],
    )

    # ══════════════════════════════════════════════════════════════
    # 8 相关文档索引
    # ══════════════════════════════════════════════════════════════
    h1(doc, "8  相关文档索引")
    add_table(
        doc,
        ["文档", "用途"],
        [
            ["skill/SKILL.md", "SOP 权威"],
            ["tools/sources.yaml", "RSS/Reddit/HN"],
            ["docs/SOURCES.md", "YAML 字段与 tag"],
            ["docs/SYSTEM_PROMPTS.md", "提示词与脚本入口"],
            ["skill/RUNBOOK.md", "逐步检查"],
            ["docs/ZenoNews_Product_Brief.docx", "对外产品简介（运行 python tools/export_product_docx.py 生成）"],
            ["web/app/about/page.tsx", "网站「如何运作」页"],
        ],
        [3.0, 4.0],
    )

    # ── Save ──
    out = Path(__file__).parent.parent / "docs" / "ZenoNews_Product_Brief.docx"
    out.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out)
    print("Saved: docs/ZenoNews_Product_Brief.docx")


if __name__ == "__main__":
    build()
