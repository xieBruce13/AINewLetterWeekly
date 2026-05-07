"""Generate the 'Newsletter Agent 说明书' Word document.

Rewrites the original Newsletter SOP into a project-reality doc:
- what this repo's multi-agent system actually does,
- the two user-facing capabilities (auto-broadcast + role-based search),
- full module structure definitions (model + product records and published sections),
- product-side first-hand source补充 (competitor social accounts + industry KOL X list).
"""

from __future__ import annotations

from pathlib import Path

from docx import Document
from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor


# ---------- small helpers ----------

NAVY = RGBColor(0x0F, 0x2E, 0x52)
AMBER = RGBColor(0xB9, 0x74, 0x00)
BLUE = RGBColor(0x1F, 0x4E, 0x9B)
MUTED = RGBColor(0x55, 0x5F, 0x6D)
BODY = RGBColor(0x1F, 0x24, 0x2B)


def _set_cell_bg(cell, hex_color: str) -> None:
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def _set_table_borders(table, color: str = "D0D7DE") -> None:
    tbl = table._tbl
    tbl_pr = tbl.find(qn("w:tblPr"))
    if tbl_pr is None:
        tbl_pr = OxmlElement("w:tblPr")
        tbl.insert(0, tbl_pr)
    borders = OxmlElement("w:tblBorders")
    for edge in ("top", "left", "bottom", "right", "insideH", "insideV"):
        e = OxmlElement(f"w:{edge}")
        e.set(qn("w:val"), "single")
        e.set(qn("w:sz"), "4")
        e.set(qn("w:color"), color)
        borders.append(e)
    tbl_pr.append(borders)


def _run(paragraph, text: str, *, bold=False, size=None, color=None, italic=False):
    r = paragraph.add_run(text)
    r.bold = bold
    r.italic = italic
    if size is not None:
        r.font.size = Pt(size)
    if color is not None:
        r.font.color.rgb = color
    r.font.name = "Microsoft YaHei"
    r._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")
    return r


def add_heading(doc, text: str, level: int) -> None:
    p = doc.add_paragraph()
    sizes = {1: 22, 2: 17, 3: 13.5, 4: 12}
    colors = {1: NAVY, 2: BLUE, 3: NAVY, 4: BODY}
    p.paragraph_format.space_before = Pt(14 if level <= 2 else 10)
    p.paragraph_format.space_after = Pt(6)
    _run(p, text, bold=True, size=sizes.get(level, 12), color=colors.get(level, BODY))


def add_para(doc, text: str, *, size=10.5, bold=False, italic=False, color=None, indent_cm=None):
    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    if indent_cm is not None:
        p.paragraph_format.left_indent = Cm(indent_cm)
    _run(p, text, bold=bold, italic=italic, size=size, color=color or BODY)
    return p


def add_bullets(doc, items: list[str], *, size=10.5, indent_cm=0.4):
    for it in items:
        p = doc.add_paragraph(style="List Bullet")
        p.paragraph_format.space_after = Pt(2)
        p.paragraph_format.left_indent = Cm(indent_cm)
        _run(p, it, size=size)


def add_number_list(doc, items: list[str], *, size=10.5):
    for it in items:
        p = doc.add_paragraph(style="List Number")
        p.paragraph_format.space_after = Pt(2)
        _run(p, it, size=size)


def add_callout(doc, title: str, body: str) -> None:
    table = doc.add_table(rows=1, cols=1)
    cell = table.cell(0, 0)
    _set_cell_bg(cell, "FFF7E6")
    cell.width = Cm(16)
    p1 = cell.paragraphs[0]
    p1.paragraph_format.space_after = Pt(2)
    _run(p1, title, bold=True, size=11, color=AMBER)
    p2 = cell.add_paragraph()
    p2.paragraph_format.space_after = Pt(0)
    _run(p2, body, size=10.5)
    _set_table_borders(table, color="E8C07D")
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_kv_table(
    doc,
    rows: list[tuple[str, str]],
    *,
    header: tuple[str, str] | None = None,
    first_col_cm: float = 3.6,
    second_col_cm: float = 12.4,
    highlight_first_row: bool = False,
) -> None:
    total_rows = len(rows) + (1 if header else 0)
    table = doc.add_table(rows=total_rows, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    _set_table_borders(table)

    r_idx = 0
    if header:
        hk, hv = header
        for col, txt in enumerate([hk, hv]):
            c = table.cell(0, col)
            _set_cell_bg(c, "0F2E52")
            c.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            _run(p, txt, bold=True, size=10.5, color=RGBColor(0xFF, 0xFF, 0xFF))
        r_idx = 1

    for i, (k, v) in enumerate(rows):
        absolute_row = r_idx + i
        c0 = table.cell(absolute_row, 0)
        c1 = table.cell(absolute_row, 1)
        if highlight_first_row and i == 0:
            _set_cell_bg(c0, "FFF7E6")
            _set_cell_bg(c1, "FFF7E6")
        else:
            _set_cell_bg(c0, "F3F6FA")
        c0.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        c1.vertical_alignment = WD_ALIGN_VERTICAL.TOP
        p0 = c0.paragraphs[0]
        p0.paragraph_format.space_after = Pt(0)
        _run(p0, k, bold=True, size=10)
        p1 = c1.paragraphs[0]
        p1.paragraph_format.space_after = Pt(0)
        _run(p1, v, size=10)

    table.columns[0].width = Cm(first_col_cm)
    table.columns[1].width = Cm(second_col_cm)
    for row in table.rows:
        row.cells[0].width = Cm(first_col_cm)
        row.cells[1].width = Cm(second_col_cm)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


def add_generic_table(doc, header: list[str], rows: list[list[str]], widths_cm: list[float]) -> None:
    table = doc.add_table(rows=len(rows) + 1, cols=len(header))
    _set_table_borders(table)
    for col, txt in enumerate(header):
        c = table.cell(0, col)
        _set_cell_bg(c, "0F2E52")
        p = c.paragraphs[0]
        p.paragraph_format.space_after = Pt(0)
        _run(p, txt, bold=True, size=10.5, color=RGBColor(0xFF, 0xFF, 0xFF))
    for i, row in enumerate(rows):
        for col, txt in enumerate(row):
            c = table.cell(i + 1, col)
            if col == 0:
                _set_cell_bg(c, "F3F6FA")
            p = c.paragraphs[0]
            p.paragraph_format.space_after = Pt(0)
            _run(p, txt, size=10, bold=(col == 0))
    for col, w in enumerate(widths_cm):
        table.columns[col].width = Cm(w)
        for row in table.rows:
            row.cells[col].width = Cm(w)
    doc.add_paragraph().paragraph_format.space_after = Pt(2)


# ---------- document content ----------

def build_doc(out_path: Path) -> None:
    doc = Document()

    # page + default style
    for section in doc.sections:
        section.top_margin = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    normal = doc.styles["Normal"]
    normal.font.name = "Microsoft YaHei"
    normal.font.size = Pt(10.5)
    normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # ---------- cover / title ----------
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.LEFT
    _run(title, "AI Newsletter Agent 说明书", bold=True, size=26, color=NAVY)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(4)
    _run(sub, "Awsome AI Newsletter · 多智能体周报流水线", size=12, color=MUTED)

    meta = doc.add_paragraph()
    meta.paragraph_format.space_after = Pt(14)
    _run(
        meta,
        "代码仓库：https://github.com/BruceXie13/AwsomeAINewsLetter ｜ 版本：重写版（对齐项目实际实现）",
        size=9.5,
        color=MUTED,
    )

    # ---------- 一、目标 ----------
    add_heading(doc, "一、项目目标", 1)
    add_para(
        doc,
        "这份周报不是 “AI 新闻搬运”。它的目标是：每周用一份结构化简报，帮团队判断过去一周里哪些模型变化、"
        "产品变化、传播变化真正值得看，并且对我们自己的产品和策略有启发。",
    )
    add_para(doc, "所以它必须同时满足三件事：", bold=True)
    add_number_list(
        doc,
        [
            "信息可信：模型以一手发布和评测为主，产品和运营要把 “热度” 与 “事实” 分开。",
            "判断清楚：不是只说发生了什么，而是要说 “为什么重要”“和我们有什么关系”。",
            "能持续复用：每周都用同一套模块、同一套过滤逻辑、同一套写作字段。",
        ],
    )

    # ---------- 二、Agent 能力总览 ----------
    add_heading(doc, "二、Agent 能支持的两项能力", 1)

    add_heading(doc, "能力 1：自动播报最新重要的 AI 资讯", 2)
    add_para(
        doc,
        "按项目内既定的 “重要性” 定义（见第四节过滤与打分），自动跑完整条流水线，产出一份结构化、"
        "可直接传阅的周报（HTML / PDF）。",
    )
    add_bullets(
        doc,
        [
            "默认时间窗口：过去 7 天；默认模块：模型 + 产品。",
            "由 10 个专职智能体分工完成（见第六节），每个 agent 只做一件事：收集 / 过滤 / 核对 / 打分 / 取舍 / 撰稿 / 质检 / 发布。",
            "产出：newsletter_runs/YYYY-MM-DD/ 目录下的 newsletter_draft.md + ai_newsletter_weekly_YYYY-MM-DD.html，必要时用 tools/render_pdf.py 导出 PDF。",
            "全过程留痕：raw / filtered / normalized / verified / scored / triage 每一步都有 JSON 中间产物，支持事后复盘与规则调参。",
        ],
    )

    add_heading(doc, "能力 2：按角色/场景支持差异化搜索", 2)
    add_para(
        doc,
        "同一套流水线，按调用方的角色、关注方向、输出深度动态切片——只要在触发时传入不同参数，"
        "即可得到不同侧重的资讯。",
    )
    add_para(doc, "每期开始前可显式指定（Scope Lock 步骤）：", bold=True)
    add_bullets(
        doc,
        [
            "time_window：时间窗口（默认最近 7 天）。",
            "modules：要跑的模块，可选 models / products / 两者都跑。",
            "top_n_per_module：每个模块最多保留几条主条目（默认 3，内容多时 5）。",
            "audience：目标读者（internal strategy / PM / investor / creator tools / general）。",
            "focus_topics：本期重点话题（例如 “coding agent”“图像生成”“设计类工具”）。",
            "must_include：必须包含的具体条目（用户指名要看的 X / Y / Z）。",
            "exclude_topics：本期排除话题。",
            "output_format：markdown / doc / pdf-ready markdown。",
        ],
    )
    add_para(doc, "典型角色切片示例：", bold=True)
    add_generic_table(
        doc,
        header=["角色", "调用时的关键参数", "输出侧重"],
        rows=[
            [
                "模型决策 / 技术选型",
                "modules=models；focus=coding/agent/长上下文；audience=internal strategy",
                "旗舰模型变化、价格/速度、对选型的直接影响",
            ],
            [
                "产品 PM / 竞品分析",
                "modules=products；focus=workflow/交互形态；must_include=重点竞品",
                "workflow 变化、交互新范式、竞品动作",
            ],
            [
                "创作者工具方向",
                "modules=products；focus=AIGC/设计/视频/素材；relevance_to_our_direction 权重拉满",
                "创作链路工具、AI 设计、视频生成、风格学习",
            ],
            [
                "投资人 / 战略视角",
                "modules=products；focus=商业模式/分发变化；top_n=5",
                "市场信号、access barrier 变化、初创聚焦",
            ],
            [
                "技术趋势监控",
                "modules=models；只保留 Tier 1+2 信源；忽略 Tier 3 newsletter",
                "官方一手 + 技术信号，不进二手解读",
            ],
            [
                "自定义主题追踪",
                "focus_topics=[任意] + must_include=[用户指定]",
                "围绕指定主题把一周内相关条目都抓来",
            ],
        ],
        widths_cm=[3.8, 6.2, 6.4],
    )
    add_para(
        doc,
        "调用方式很轻量：在 Claude Code / Cursor 里说一句中文指令即可，例如 “按 newsletter SOP 跑本周周报，"
        "只看产品模块，重点关注 creator tools 方向”。",
        italic=True,
        color=MUTED,
        size=10,
    )

    # ---------- 三、产出资讯的结构化定义 ----------
    add_heading(doc, "三、产出资讯包含哪些内容（模块结构定义）", 1)
    add_para(
        doc,
        "下面给每个模块的 “记录结构（record）” 和 “周报成品结构（output）” 一一定义。"
        "记录结构是过滤/打分/取舍的中间态；成品结构是读者最终看到的内容。",
    )

    # 3.1 模型模块 record
    add_heading(doc, "3.1 模型模块 · 条目记录结构（model_record）", 2)
    add_kv_table(
        doc,
        header=("字段组", "字段说明"),
        rows=[
            ("基础信息", "module=model ｜ name ｜ version ｜ company ｜ updated_at"),
            (
                "信源字段",
                "source_tier（official / third-party-validation / secondary / community-only）"
                "｜source_primary[]（一手链接）｜source_validation[]（Arena / Artificial Analysis / LiveBench）"
                "｜source_secondary[]（newsletter / 媒体）｜raw_urls[]",
            ),
            (
                "内容字段（官方）",
                "update_type（新旗舰/版本升级/能力跳跃/价格变更/API 能力/开源权重/训练方法…）"
                "｜official_claims[]（厂商原话 + 数字）",
            ),
            (
                "内容字段（第三方 & 编辑）",
                "external_validation_summary ｜ real_change_notes ｜ selection_impact_notes"
                "｜price_speed_cost_notes ｜ app_api_workflow_notes ｜ new_use_cases[]"
                "｜business_model ｜ ecosystem_echo ｜ durability_notes ｜ risk_caveat",
            ),
            (
                "社区反馈",
                "user_market_feedback.good[] ｜ user_market_feedback.bad[]（来自 Reddit / HN / X 真实用户原话）",
            ),
            (
                "判断与关系",
                "one_line_judgment（一句话判断） ｜ relevance_to_us（与我们产品的相关性）",
            ),
            (
                "过滤结果（gate_results）",
                "is_model_event ｜ has_tier1_or_tier2_source ｜ minimum_evidence_met",
            ),
            (
                "打分结果（score_breakdown）",
                "real_capability_change ｜ selection_impact ｜ evidence_quality ｜ ecosystem_echo"
                "｜durability ｜ hype_penalty ｜ total ｜ justifications{}",
            ),
            (
                "取舍",
                "verification_status（verified / partially / unverified） ｜ confidence（high / medium / low）"
                "｜item_tier（main / brief / dropped） ｜ drop_reason",
            ),
            (
                "配图",
                "image_urls[]：Step 1 收集时顺手存好的真实产品 / 跑分截图 URL（禁用 hero banner / 营销图）",
            ),
        ],
        first_col_cm=4.0,
        second_col_cm=12.0,
    )

    # 3.2 产品模块 record
    add_heading(doc, "3.2 产品模块 · 条目记录结构（product_record）", 2)
    add_kv_table(
        doc,
        header=("字段组", "字段说明"),
        rows=[
            ("基础信息", "module=product ｜ name ｜ company ｜ updated_at"),
            (
                "信源字段",
                "source_tier（official-changelog / discovery-platform / media / community-only）"
                "｜source_primary[] ｜ source_secondary[] ｜ raw_urls[]",
            ),
            (
                "定位与问题",
                "launch_signal（冷启动/公测/GA/灰度等） ｜ core_positioning（一句话定位）"
                "｜problem_it_solves（解决什么问题）",
            ),
            (
                "变化维度",
                "user_visibility（用户能否今天就感知到） ｜ access_barrier_change（免费/新地区/API 开放等）"
                "｜workflow_change（工作流是否被真正改变） ｜ distribution_change（分发/系统级集成变化）"
                "｜new_product_form（新的产品形态，如桌宠/常驻代理/语音前端/工作流中心）"
                "｜interaction_pattern（交互范式变化） ｜ user_scenarios[]",
            ),
            ("商业模式", "business_model（定价 + 包装策略） ｜ risk_caveat"),
            (
                "市场反馈",
                "market_feedback.good[] ｜ market_feedback.bad[] ｜ market_signal_strength"
                "（Product Hunt 名次 / 用户自发讨论 / 稳定用户案例）",
            ),
            (
                "判断与关系",
                "one_line_judgment ｜ relevance_to_us（与我们创作/AIGC/设计方向的直接相关性）",
            ),
            (
                "过滤结果（gate_results）",
                "user_can_perceive_today ｜ evidence_beyond_single_source",
            ),
            (
                "打分结果（score_breakdown）",
                "user_visibility ｜ access_barrier_change ｜ workflow_change ｜ distribution_change"
                "｜user_reaction ｜ relevance_to_our_direction ｜ evidence_quality ｜ hype_penalty"
                "｜total ｜ justifications{}",
            ),
            (
                "取舍",
                "verification_status ｜ confidence ｜ item_tier（main / brief / dropped） ｜ drop_reason",
            ),
            (
                "配图",
                "image_urls[]：真实 UI 截图 / 竞品 / 对比表；严禁 hero banner / 概念插画 / Logo 拼图",
            ),
        ],
        first_col_cm=4.0,
        second_col_cm=12.0,
    )

    # 3.3 成品结构
    add_heading(doc, "3.3 周报成品结构（读者看到的版面）", 2)
    add_para(doc, "每期周报按固定 7 段结构呈现，总篇幅目标 6–8 页 PDF：", bold=True)
    add_generic_table(
        doc,
        header=["段落", "内容 & 写法"],
        rows=[
            ["① 本周结论", "1–2 个琥珀色 callout block，直接说本期核心判断 + 对我们意味着什么。"],
            [
                "② 模型模块",
                "1 句模块总结 + 每条主条目一张卡片：总结 / 模型能力 / 产品功能 / 新使用场景 / 商业模式 / 用户反馈 / 与我们的关系。",
            ],
            [
                "③ 产品模块",
                "1 句模块总结 + 每条主条目一张卡片：总结 / 核心定位 / 产品重点 / 用户场景 / 商业模式 / 用户反馈 / 与我们的关系。",
            ],
            [
                "④ 初创聚焦",
                "每期强制至少 1 个非大厂的初创/独立产品（主栏或简讯出现），字段：总结 / 核心机制 / 典型场景 / 为什么关注。",
            ],
            ["⑤ 简讯", "一张紧凑 2 列表格：名称 + 一句话。不写分数、不写理由，只作扫描用。"],
            ["⑥ 编辑部判断", "趋势（2 条） + 项目动作（可执行） + 下周监控。"],
            ["⑦ 参考来源", "按模型 / 产品分组，最小号灰字，便于溯源但不挤占版面。"],
        ],
        widths_cm=[3.2, 13.0],
    )
    add_para(doc, "每张主条目卡片都必须包含的 “信息分层 2×2 小表”：", bold=True)
    add_generic_table(
        doc,
        header=["官方声明", "外部验证"],
        rows=[["厂商原话 + 数字，可点链", "Arena / Artificial Analysis / LiveBench 等"]],
        widths_cm=[8.0, 8.0],
    )
    add_generic_table(
        doc,
        header=["社区反馈", "编辑判断"],
        rows=[["2–4 条 Reddit 真实用户原话（全部译为中文）", "我们怎么看 + 与我们的关系"]],
        widths_cm=[8.0, 8.0],
    )

    # ---------- 四、过滤与打分逻辑 ----------
    add_heading(doc, "四、过滤与打分逻辑（什么叫 “重要”）", 1)
    add_para(doc, "门槛和权重由 skill/rubric.json 统一定义，每一条都留存 justification。", color=MUTED, size=10)

    add_heading(doc, "4.1 模型模块", 2)
    add_para(doc, "门槛（未过则丢弃或进观察池）：", bold=True)
    add_bullets(
        doc,
        [
            "is_model_event：必须是模型事件（新模型 / 版本升级 / 能力跳跃 / 价格/速度/上下文重大变化 / API 能力 / 开源权重 / 重要训练方法）。公司融资、高管发言、一般 AI 评论直接丢弃。",
            "has_tier1_or_tier2_source：至少要有一个官方一手或技术信号一手（HF / GitHub Trending / HN / arXiv / 基准库）；只在二手 newsletter 里出现的进观察池。",
            "minimum_evidence_met：至少一条可核对的 claim + 可追溯的 URL。",
        ],
    )
    add_para(doc, "打分维度（总分理论值 -3 ~ 10）：", bold=True)
    add_generic_table(
        doc,
        header=["维度", "区间", "简述"],
        rows=[
            ["real_capability_change", "0–3", "能力是不是真的变了（reasoning / coding / agent / multimodal / context）"],
            ["selection_impact", "0–2", "会不会改变产品团队的模型选型"],
            ["evidence_quality", "0–2", "官方 + 独立第三方验证才算满分"],
            ["ecosystem_echo", "0–2", "生态是否已经在适配 / 搭建"],
            ["durability", "0–1", "能撑 1 个月以上才给分"],
            ["hype_penalty", "−3–0", "纯公关稿/demo 热度减分"],
        ],
        widths_cm=[4.2, 2.0, 9.8],
    )
    add_para(doc, "分档阈值：main ≥ 7 ｜ brief 5–6 ｜ drop < 5。编辑 override：两条编辑门（是否真变化 + 是否影响选型）任一不过则降级。", bold=True)

    add_heading(doc, "4.2 产品模块", 2)
    add_para(doc, "门槛：", bold=True)
    add_bullets(
        doc,
        [
            "user_can_perceive_today：公开可用 / 公测 / 明确将发。纯概念或 closed alpha 最多进简讯。",
            "evidence_beyond_single_source：至少两个独立来源；单个 newsletter 孤证进观察池。",
        ],
    )
    add_para(doc, "打分维度（总分理论值 -3 ~ 16）：", bold=True)
    add_generic_table(
        doc,
        header=["维度", "区间", "简述"],
        rows=[
            ["user_visibility", "0–2", "用户今天能不能真感知到"],
            ["access_barrier_change", "0–2", "免费/新地区/新平台/API 开放等准入变化"],
            ["workflow_change", "0–3", "任务完成方式是否被根本改变（单轮→多步代理，文本→多模态…）"],
            ["distribution_change", "0–2", "新平台 / 系统级集成 / 重大合作"],
            ["user_reaction", "0–2", "Reddit / X / HN 是否自发讨论、是否有真实用户原话"],
            ["relevance_to_our_direction", "0–3", "创作 / AIGC / 设计 / 视频 / 素材 / 风格相关的权重更高"],
            ["evidence_quality", "0–2", "官方 + 市场信号（PH 名次、用户数、社区采用）"],
            ["hype_penalty", "−3–0", "Vaporware / 演示向减分"],
        ],
        widths_cm=[4.2, 2.0, 9.8],
    )
    add_para(doc, "分档阈值：main ≥ 10 ｜ brief 7–9 ｜ drop < 7。", bold=True)
    add_callout(
        doc,
        "两条结构性规则",
        "① 多样性：Top N 主条目里同公司最多 N−1 条，超出就把最弱的那条降级到简讯。"
        "② 初创必带：每期（含简讯）至少一个非大厂（不属于 OpenAI / Anthropic / Google / Meta / Microsoft / Adobe / Apple）的独立/初创产品。",
    )

    # ---------- 五、信息源总览（含产品侧一手补充） ----------
    add_heading(doc, "五、信息源总览（产品侧 todo 已补充一手账号）", 1)

    add_heading(doc, "5.1 模型模块信源", 2)
    add_generic_table(
        doc,
        header=["分层", "典型来源"],
        rows=[
            [
                "Tier 1 一手（优先）",
                "OpenAI News + Release Notes、Anthropic Newsroom + Claude Release Notes、Google Gemini 官方 + Gemini Apps Release Notes、Meta AI Blog、Mistral、xAI News + Dev Release Notes。"
                "次要：GLM、MiniMax、Kimi、Midjourney、Seedream、Flux、其他新兴实验室。",
            ],
            [
                "Tier 2 技术信号",
                "Hugging Face Trending Papers / Models、GitHub Trending、Hacker News、arXiv discovery feeds、主要基准库新 eval run。",
            ],
            [
                "外部验证",
                "Arena / LMSYS 类榜单、Artificial Analysis、LiveBench。",
            ],
            [
                "Tier 3 二手（补漏）",
                "Import AI、TLDR AI、The Rundown AI。",
            ],
            [
                "Prompt 补充",
                "用户一句 “这周看一下 xxx” 即算强制收集目标。",
            ],
        ],
        widths_cm=[3.8, 12.2],
    )

    add_heading(doc, "5.2 产品模块信源", 2)
    add_para(doc, "① 官方一手 changelog（Tier 1，必扫）", bold=True)
    add_bullets(
        doc,
        [
            "ChatGPT Release Notes + Enterprise/Edu Release Notes",
            "Claude Release Notes",
            "Gemini Apps Release Notes + 每月 Gemini Drops",
            "Perplexity Changelog",
            "xAI Developer Release Notes",
        ],
    )

    add_para(doc, "② 重点竞品官方社媒账号（补充：一手信号的第二条通道）", bold=True)
    add_para(
        doc,
        "产品更新常常先在 X/YouTube/TikTok 官方账号冒头，再出 changelog。按这类账号建定时抓取，"
        "可显著领先二手 newsletter。",
        color=MUTED,
        size=10,
    )
    add_generic_table(
        doc,
        header=["类别", "官方账号（X / YouTube / TikTok）"],
        rows=[
            [
                "基础模型厂商",
                "@OpenAI ｜ @AnthropicAI ｜ @GoogleDeepMind ｜ @GeminiApp ｜ @MetaAI ｜ @MistralAI ｜ @xai ｜ @AIatMeta",
            ],
            [
                "面向用户的 AI 产品",
                "@perplexity_ai ｜ @ChatGPTapp ｜ @claudeai ｜ @character_ai ｜ @poe_platform",
            ],
            [
                "创作 / 图像 / 视频 / 音频",
                "@midjourney ｜ @runwayml ｜ @pika_labs ｜ @LumaLabsAI ｜ @suno_ai_ ｜ @udiomusic ｜ @elevenlabsio ｜ @heygen_official",
            ],
            [
                "设计 / 创作工具",
                "@figma ｜ @canva ｜ @Adobe ｜ @framer ｜ @gamma_app ｜ @notion ｜ @linear",
            ],
            [
                "Agent / 编码工具",
                "@cursor_ai ｜ @replit ｜ @v0 ｜ @lovable_dev ｜ @bolt_new ｜ @devin_ai ｜ @windsurf_ai ｜ @githubcopilot",
            ],
            [
                "硬件 / 端侧",
                "@Humane ｜ @rabbit_hmi ｜ @MetaAI（Ray-Ban）",
            ],
        ],
        widths_cm=[3.8, 12.2],
    )

    add_para(doc, "③ 行业重要 KOL 的 X / Twitter（补充：真实用户+行业判断的一手通道）", bold=True)
    add_generic_table(
        doc,
        header=["类别", "KOL 账号"],
        rows=[
            [
                "AI 公司创始人 / 核心高管",
                "@sama（Sam Altman）｜ @gdb（Greg Brockman）｜ @miramurati ｜ @DarioAmodei ｜ @darioamodei ｜ "
                "@demishassabis ｜ @JeffDean ｜ @sundarpichai ｜ @elonmusk ｜ @AravSrinivas ｜ @mustafasuleyman ｜ @alexandr_wang",
            ],
            [
                "研究 / 技术布道",
                "@karpathy ｜ @ylecun ｜ @ilyasut ｜ @emollick ｜ @alexalbert__ ｜ @OfficialLoganK ｜ "
                "@_philschmid ｜ @drjimfan ｜ @natolambert ｜ @jeremyphoward",
            ],
            [
                "AI 工程 / 评测",
                "@swyx ｜ @simonw ｜ @goodside ｜ @marktenenholtz ｜ @hwchase17（LangChain） ｜ @jerryjliu0（LlamaIndex） ｜ @abacaj",
            ],
            [
                "产品 / 市场 / 投资",
                "@pmarca（Marc Andreessen）｜ @levie（Aaron Levie）｜ @bgurley ｜ @gabemkahn ｜ @packyM ｜ @lennysan",
            ],
            [
                "创作 / 设计 / 视频方向",
                "@zoink（Dylan Field，Figma）｜ @bilawalsidhu ｜ @nickfloats ｜ @cristobalvalenzu ｜ @blizaine ｜ @heyBarsee ｜ @icreatelife",
            ],
            [
                "中文圈（可选）",
                "@dotey ｜ @op7418 ｜ @GoodIdeaAI ｜ @oran_ge ｜ @jikeshijian",
            ],
        ],
        widths_cm=[3.8, 12.2],
    )
    add_para(
        doc,
        "说明：以上账号作为产品模块 Tier 1 的 “一手信号补充”，与官方 changelog 同级扫描；"
        "但内容本身仍需满足 Gate Check（用户今天可感知 + 至少两个独立来源）才能进主栏。",
        italic=True,
        color=MUTED,
        size=10,
    )

    add_para(doc, "④ 发现平台 & 分发信号（Tier 2）", bold=True)
    add_bullets(
        doc,
        [
            "发现：Product Hunt（AI 专区）、Futurepedia、There's An AI For That、Toolify。",
            "分发：Google Trends、TikTok Creative Center、YouTube Charts、X Trends。",
        ],
    )
    add_para(doc, "⑤ Newsletter / 媒体（Tier 3，用于补漏与解释）", bold=True)
    add_bullets(
        doc,
        [
            "The Rundown AI、Ben's Bites、Superhuman AI、TLDR AI、AI Breakfast、AI Valley。",
            "The Verge (AI)、WIRED (AI)、TechCrunch (AI)、AI News、THE DECODER、VentureBeat (AI)。",
        ],
    )
    add_para(doc, "⑥ 社区 / 生态（Tier 2–3）", bold=True)
    add_bullets(
        doc,
        [
            "Reddit：r/artificial、r/singularity、r/LocalLLaMA、r/ChatGPT、r/ClaudeAI、r/MachineLearning、r/SideProject、r/startups。",
            "Hacker News、YC Startup Directory (AI)、Hugging Face Hub、Towards AI、Indie Hackers。",
        ],
    )

    # ---------- 六、Agent 流水线 ----------
    add_heading(doc, "六、Agent 流水线（每个 agent 具体干什么）", 1)
    add_para(
        doc,
        "项目在 .claude/agents/ 下有 10 个专职智能体；由 orchestrator 协调，每个 agent 只读自己需要的 SOP 片段，"
        "保持上下文最小、行为稳定。",
    )
    add_generic_table(
        doc,
        header=["Step", "Agent", "读取", "写出"],
        rows=[
            ["0", "orchestrator", "用户请求 + SOP", "run_header.md（本期范围锁定）"],
            ["1", "collector", "run_header + §1A / §1B 信源列表", "raw_model_records.json / raw_product_records.json + 真实产品图片 URL"],
            ["2", "filter", "raw + rubric.json 门槛", "filtered_records.json（未过 gate 的进 watchlist 或 drop）"],
            ["3", "normalizer", "filtered + record_schemas.json", "normalized_records.json（扩成完整 schema）"],
            ["4 + 4B", "verifier", "normalized", "verified_records.json（含 Reddit 真实用户 2–4 条原话）"],
            ["5", "scorer", "verified + rubric 维度", "scored_records.json（每维都带 justification）"],
            ["6", "triage-editor", "scored + 阈值 + 多样性/初创规则", "triage_decisions.json（main / brief / drop）"],
            ["7", "writer", "triage + verified + output_template.md", "newsletter_draft.md（用户原话翻成中文）"],
            ["8", "qa-reviewer", "draft + 清单", "PASS 或 REVISE 意见"],
            ["9", "publisher", "draft + tools/convert_to_pdf.py", "ai_newsletter_weekly_YYYY-MM-DD.html（必要时导出 PDF）"],
        ],
        widths_cm=[1.3, 3.2, 5.5, 6.0],
    )
    add_para(
        doc,
        "orchestrator 本身不写新闻内容；所有规则写在 skill/SKILL.md 里作为唯一真源，agent prompt 都引用它。",
        italic=True,
        color=MUTED,
        size=10,
    )

    # ---------- 七、使用方式 ----------
    add_heading(doc, "七、使用方式", 1)
    add_number_list(
        doc,
        [
            "安装本地渲染依赖：pip install -r requirements.txt（PDF 导出另装 requirements-pdf.txt + playwright install chromium）。",
            "在智能体里一句话触发：例如 “按 newsletter SOP 跑本周周报” 或带参数 “只看产品模块，重点关注 workflow 变化和创作方向相关性”。",
            "到 newsletter_runs/<本期日期>/ 取稿：主文件是 newsletter_draft.md，同目录下有素材图和网页文件。",
            "需要 PDF：进入该期目录执行 python ../../tools/render_pdf.py，得到与网页一致版式的 PDF。",
        ],
    )
    add_heading(doc, "每期归档产物", 2)
    add_bullets(
        doc,
        [
            "raw_model_records.json / raw_product_records.json",
            "filtered_records.json / normalized_records.json / verified_records.json",
            "scored_records.json / triage_decisions.json",
            "newsletter_draft.md",
            "images/（当期真实产品截图 & 跑分图表）",
            "ai_newsletter_weekly_YYYY-MM-DD.html（+ 可选 PDF）",
        ],
    )

    # ---------- 八、编辑原则 ----------
    add_heading(doc, "八、编辑原则（与读者相关的硬规则）", 1)
    add_bullets(
        doc,
        [
            "官方声明 / 外部验证 / 社区反馈 / 编辑判断 四层分明，绝不混在一段里。",
            "每条主条目 2–4 条真实用户原话，全部翻成中文；泛夸不算。",
            "只用真实产品 UI / 跑分截图；hero banner / 概念插画 / AI 生成艺术 / 缩略图一律不用——宁可不配图。",
            "正文不暴露分数、置信度、Tier 标签——那些是流程产物，不是给读者看的。",
            "句子短、先下结论再给证据；关键数字和判断词加粗，让读者扫粗体就能拿 80% 内容。",
            "每期版面目标 6–8 页 PDF；保证扫描路径：标题 → 本周结论 → 每张卡片的 “总结” 行。",
        ],
    )

    doc.save(out_path)
    print(f"Wrote docx ({out_path.stat().st_size} bytes)")


if __name__ == "__main__":
    out = Path(__file__).resolve().parent.parent / "Newsletter_Agent_说明书.docx"
    build_doc(out)
